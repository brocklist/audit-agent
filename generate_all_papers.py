"""
批量生成16科目审计底稿 — 离线版
使用 audit_engine 确定性计算 + LLM判断
"""
import os, sys, json, time, asyncio
from pathlib import Path
from docx import Document
from docx.shared import Pt
from datetime import datetime

# 确保能找到audit_engine
sys.path.insert(0, str(Path(__file__).parent))

from audit_engine import (
    CALCULATORS, SUBJECT_NAMES, load_subject_data,
    call_llm, extract_json, ExcelRenderer, OUTPUT_DIR, ENTITY, PERIOD, AUDITOR, FIRM
)

OUTPUT_DIR.mkdir(exist_ok=True)


def generate_word_log(subject_id, subject_name, calc, llm_result, elapsed, xlsx_path):
    """生成Word格式审计日志"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs: run.font.name = '微软雅黑'

    def add_table(headers, rows):
        table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Grid Accent 1')
        for i, h in enumerate(headers): table.rows[0].cells[i].text = str(h)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row): table.rows[ri+1].cells[ci].text = str(val)
        doc.add_paragraph()

    add_heading(f'审计日志 - {subject_name}（{subject_id}）', 0)

    # 基本信息
    add_heading('审计基本信息', 1)
    add_table(['项目', '内容'], [
        ['被审计单位', ENTITY],
        ['审计科目', f'{subject_name}（索引号:{subject_id}）'],
        ['会计期间', PERIOD],
        ['审核员', AUDITOR],
        ['事务所', FIRM],
        ['审计日期', datetime.now().strftime('%Y-%m-%d')],
        ['总耗时', f'{elapsed:.1f}秒'],
    ])

    # 审计发现
    add_heading('审计发现', 1)
    findings = llm_result.get("findings", [])
    for i, f in enumerate(findings, 1):
        doc.add_paragraph(f"{i}. {f}", style='List Number')

    # 风险清单
    risks = llm_result.get("risk_items", [])
    if risks:
        add_heading('风险清单', 1)
        add_table(['风险等级', '项目', '金额', '原因'],
                  [[r.get('level',''), r.get('item',''), r.get('amount',''), r.get('reason','')] for r in risks])

    # 各sheet结论
    add_heading('底稿Sheet审计结论', 1)
    conclusions = llm_result.get("conclusions", {})
    for sheet_title, conclusion in conclusions.items():
        doc.add_paragraph(f"{sheet_title}: {conclusion}")

    # 整体摘要
    add_heading('整体审计摘要', 1)
    doc.add_paragraph(llm_result.get("summary", ""))

    # 技术信息
    add_heading('技术信息', 2)
    add_table(['项目', '内容'], [
        ['生成方式', '代码计算 + LLM判断'],
        ['计算Sheet数', str(len(calc.sheets))],
        ['输出文件', xlsx_path],
    ])

    # 保存
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = f"{subject_id}_{subject_name}_审计日志.docx"
    path = OUTPUT_DIR / fname
    doc.save(str(path))
    return str(path)


async def generate_one(subject_id):
    """为单个科目生成审计底稿"""
    subject_name = SUBJECT_NAMES[subject_id]
    print(f"\n{'='*60}")
    print(f"[{subject_id}] {subject_name} 开始生成...")
    print(f"{'='*60}")

    total_start = time.time()

    # Step 1: 加载数据
    print(f"  [1/5] 加载案例数据...")
    data = load_subject_data(subject_id)
    if not data:
        print(f"  [SKIP] 无数据文件")
        return False
    print(f"  文件数: {len(data)}")

    # Step 2: 执行确定性计算
    print(f"  [2/5] 执行计算...")
    calc_class = CALCULATORS[subject_id]
    calc = calc_class(data)
    calc.calculate()
    print(f"  生成 {len(calc.sheets)} 张Sheet")

    # Step 3: 调用LLM获取判断
    print(f"  [3/5] LLM判断...")
    prompt = calc.build_llm_prompt()
    prompt_size = len(prompt)
    print(f"  Prompt: {prompt_size:,} 字符")

    llm_start = time.time()
    try:
        llm_response = await call_llm(prompt)
        llm_elapsed = time.time() - llm_start
        print(f"  LLM响应: {len(llm_response):,} 字符 · {llm_elapsed:.0f}s")
    except Exception as e:
        print(f"  [ERR] LLM调用失败: {e}")
        # 使用空结果继续
        llm_response = '{"summary":"LLM调用失败","risk_items":[],"findings":[],"conclusions":{}}'

    # Step 4: 解析LLM结果 + 合并结论
    print(f"  [4/5] 解析结果...")
    llm_result = extract_json(llm_response)
    calc.merge_conclusions(llm_result)

    # Step 5: 渲染Excel + Word
    print(f"  [5/5] 渲染输出...")
    renderer = ExcelRenderer(ENTITY, PERIOD, AUDITOR, FIRM, subject_id)
    for s in calc.sheets:
        renderer.add_sheet(
            title=s.get("title", "Sheet"),
            headers=s.get("headers", []),
            rows=s.get("rows", []),
            col_widths=s.get("col_widths", [10]*len(s.get("headers",[]))),
            conclusion=s.get("conclusion", ""),
            col_types=s.get("col_types", None),
            highlights=s.get("highlights", None),
            sheet_no=s.get("sheet_no", "")
        )
    xlsx_path, alt_path = renderer.save(subject_name)
    print(f"  Excel: {alt_path}")

    elapsed = time.time() - total_start
    log_path = generate_word_log(subject_id, subject_name, calc, llm_result, elapsed, xlsx_path)
    print(f"  Log: {log_path}")

    # Summary
    summary_text = llm_result.get("summary", "")
    risk_count = len(llm_result.get("risk_items", []))
    finding_count = len(llm_result.get("findings", []))

    print(f"  [OK] {subject_name} 完成 · {elapsed:.0f}s")
    print(f"  风险: {risk_count} · 发现: {finding_count}")
    print(f"  摘要: {summary_text[:80]}...")

    return {
        "subject": subject_id,
        "subject_name": subject_name,
        "success": True,
        "xlsx": xlsx_path,
        "alt_xlsx": alt_path,
        "log": log_path,
        "sheets": [s["title"] for s in calc.sheets],
        "summary": summary_text,
        "risk_items": llm_result.get("risk_items", []),
        "findings": llm_result.get("findings", []),
        "elapsed": round(elapsed, 1),
    }


async def main():
    import argparse
    parser = argparse.ArgumentParser(description='批量生成审计底稿')
    parser.add_argument('subjects', nargs='*', default=[],
                       help='科目代码，如 C D E ... (默认全部)')
    parser.add_argument('--parallel', action='store_true', help='并行生成')
    args = parser.parse_args()

    if args.subjects:
        subject_ids = [s.upper() for s in args.subjects if s.upper() in CALCULATORS]
    else:
        subject_ids = list(CALCULATORS.keys())

    print(f"审计底稿生成器 v2")
    print(f"被审计单位: {ENTITY}")
    print(f"会计期间: {PERIOD}")
    print(f"科目数: {len(subject_ids)}")
    print(f"输出目录: {OUTPUT_DIR}")

    results = []

    if args.parallel:
        # 并行模式（小心API限流）
        tasks = [generate_one(sid) for sid in subject_ids]
        results = await asyncio.gather(*tasks)
    else:
        # 顺序模式
        for sid in subject_ids:
            result = await generate_one(sid)
            if result:
                results.append(result)
            # 短暂休息避免API限流
            await asyncio.sleep(2)

    # 汇总
    print(f"\n{'='*60}")
    print(f"ALL DONE! 成功: {len(results)}/{len(subject_ids)}")
    print(f"{'='*60}")

    for r in results:
        if r:
            print(f"  [{r['subject']}] {r['subject_name']}: {r['elapsed']:.0f}s · {len(r['sheets'])} sheets · {len(r['risk_items'])} risks")

    # 列出输出文件
    print(f"\n输出文件:")
    for f in sorted(OUTPUT_DIR.glob("*")):
        if f.is_file():
            sz = f.stat().st_size
            print(f"  {f.name} ({sz:,} bytes)")


if __name__ == "__main__":
    asyncio.run(main())
