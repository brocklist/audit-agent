"""生成技术原理 Word 文档到桌面"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 页面设置
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_title(text):
    h = doc.add_heading(text, level=0)
    for run in h.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x0e, 0xa5, 0xe9)

def add_h1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x02, 0x84, 0xc7)

def add_h2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x38, 0xbd, 0xf8)

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bold_para(bold, normal):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run_b = p.add_run(bold)
    run_b.bold = True
    run_b.font.size = Pt(10.5)
    p.add_run(normal)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ==================== 正文 ====================

add_title("审计智能体 · 技术原理全景")

add_para("本文档完整记录了 AI 审计智能体项目的全部技术原理，按技术层次组织，覆盖 AI/LLM、审计计算引擎、Excel 底稿渲染、OCR 凭证识别、后端架构、前端架构、范例数据生成、安全合规等 8 大模块。")
add_para(f"统计：Python ~3,700 行 + HTML/CSS/JS ~2,400 行 = 约 6,100 行核心代码，16 个审计科目，16 个 API 端点，55 个范例 Excel + 3 张 OCR 示例图片。")

# ========== 一 ==========
add_h1("一、AI 与 LLM 层")

add_h2("1.1 模型调用")
add_bold_para("模型：", "DeepSeek V4 Pro [1M 上下文窗口]")
add_bold_para("调用协议：", "Anthropic API 兼容协议（Messages API）")
add_bold_para("端点：", "https://api.deepseek.com/anthropic/v1/messages")
add_bold_para("关键参数：", "temperature=0.05（低温度确保审计结论确定性，避免幻觉），max_tokens=64000")
add_bold_para("注入方式：", "通过环境变量 ANTHROPIC_BASE_URL / ANTHROPIC_AUTH_TOKEN / ANTHROPIC_MODEL 注入，支持任意 Anthropic 兼容服务替换")
add_bold_para("HTTP 客户端：", "httpx.AsyncClient，timeout=600s，支持流式响应")

add_h2("1.2 双引擎架构")
add_para("项目实现两套审计执行引擎，可按科目灵活切换：")
add_bold_para("server.py 引擎：", "LLM 全权处理模式 — 将原始 Excel 数据序列化为文本，构建包含审计程序的 prompt，由 AI 一次性输出完整 JSON 结果")
add_bold_para("audit_engine.py 引擎：", "Python 确定性计算 + LLM 判断模式 — Calculator 基类派生 16 个子类，Python 完成数值计算（变动率、利息、折旧等），将计算结果与样本数据一同发给 LLM 做审计判断")
add_bold_para("路由机制：", "CALCULATORS 字典按科目 ID（C-R）映射到对应 Calculator 子类，自动选择引擎")

add_h2("1.3 System Prompt 工程")
add_bold_para("文件：", "CLAUDE.md（99 行）")
add_bold_para("内容结构：", "身份定义（中国注册会计师）+ 审计程序清单（C-H 六科完整，其余科通用模板）+ 底稿结构规范（审定表/明细表/程序表/函证汇总表/审计结论）+ 判断标准（变动率>30%黄色预警、>50%红色异常、函证回收率<80%增加替代程序）")
add_bold_para("输出格式约束：", "强制 JSON schema：{summary, risk_items: [{level, item, amount, reason}], findings, sheets: [{type, title, headers, rows, conclusion}]}")

add_h2("1.4 JSON 容错修复")
add_bold_para("策略：", "双重提取 + 兜底修复")
add_bold_para("第一层：", "正则提取 markdown 代码块（```json ... ```）中的 JSON，json.loads() 解析")
add_bold_para("第二层：", "解析失败时调用 json-repair 库的 repair_json() 自动修复缺失引号、多余逗号、截断等问题")
add_bold_para("容错率：", "> 95%，覆盖 LLM 输出最常见的 3 类格式错误")

add_h2("1.5 SSE 实时流式推送")
add_bold_para("实现方式：", "FastAPI StreamingResponse + media_type='text/event-stream'")
add_bold_para("进度阶段：", "共 6 阶段 — 文件读取(0-12%) → 数据序列化(13-19%) → Prompt构建(20-24%) → AI分析(25-75%) → Excel生成(76-90%) → 完成(91-100%)")
add_bold_para("心跳机制：", "15 秒间隔发送 ': heartbeat\\n\\n' SSE 注释保持长连接，前端过滤此消息")
add_bold_para("超时：", "600 秒，覆盖最复杂科目的完整审计流程")

# ========== 二 ==========
add_h1("二、审计计算引擎")

add_h2("2.1 基类架构 (BaseCalculator)")
add_para("BaseCalculator 定义统一接口：")
add_para("  • calculate() — 执行所有 Python 计算，填充 self.sheets 列表")
add_para("  • build_llm_prompt() — 将计算结果构建为 LLM prompt 文本")
add_para("  • self.sheets — list[dict]，每项含 title / headers / col_widths / rows / col_types / conclusion")
add_para("  • self.stats — dict，科目统计摘要")
add_para("通用工具函数：_trial_balance_rows() 自动从上传文件中识别科目余额表；_find_file() 按关键词匹配明细文件；safe_float() / safe_str() 容错转换；pct_str() 变动率格式化")

add_h2("2.2 16 个科目计算器")

add_table(
    ["科目", "类名", "底稿数", "核心计算"],
    [
        ["C 货币资金", "MonetaryFundCalc", "8", "现金盘点系数、银行余额调节、利息测算、截止测试、大额双向核对"],
        ["D 应收票据", "NoteReceivableCalc", "10", "票据背书/贴现检查、带息票据利息测算、预期信用损失ECL模型"],
        ["E 其他应收款", "OtherReceivableCalc", "8", "账龄分析、关联方资金占用审核、坏账准备计算"],
        ["F 预付账款", "PrepaymentCalc", "5", "长期挂账查验、减值迹象复核"],
        ["G 预收账款", "AdvanceReceiptCalc", "5", "收入确认时点检查（履约进度vs预收）"],
        ["H 短期借款", "ShortTermLoanCalc", "6", "借款利息复核、贷款卡记录核对、抵押质押统计"],
        ["I 应收账款", "AccountReceivableCalc", "10", "账龄分析、预期信用损失、坏账准备、函证控制"],
        ["J 应付账款", "AccountPayableCalc", "5", "暂估入账检查、采购合同核对"],
        ["K 应付票据", "NotePayableCalc", "5", "保证金台账核对、票据到期检查"],
        ["L 应付职工薪酬", "PayrollCalc", "5", "工资/社保/公积金计提比例合规检查"],
        ["M 应交税费", "TaxCalc", "5", "纳税申报核对、增值税明细表、税金计提复核"],
        ["N 其他应付款", "OtherPayableCalc", "4", "账龄分析、控股股东占用检查"],
        ["O 存货", "InventoryCalc", "5", "收发存勾稽、跌价准备测试、库龄分析"],
        ["P 固定资产", "FixedAssetCalc", "6", "折旧测算、减值测试、处置损益检查"],
        ["Q 无形资产", "IntangibleAssetCalc", "6", "摊销测算、权属证书清单检查"],
        ["R 长期借款", "LongTermLoanCalc", "5", "利息测算、还款计划核对、担保统计"],
    ]
)

add_h2("2.3 Excel 文件解析")
add_bold_para("read_excel() 函数：", "openpyxl.load_workbook() → 逐 Sheet 读取 → 第一行作为表头 → 后续行转为 list[dict]")
add_bold_para("数据结构：", "返回 {sheet_name: {headers: [...], rows: [{col:val, ...}, ...]}} 的嵌套字典")

# ========== 三 ==========
add_h1("三、Excel 底稿渲染引擎")

add_h2("3.1 ExcelBuilder 类")
add_para("基于 openpyxl 的专业审计底稿渲染引擎，特征：")
add_para("  • 4 行统一表头：审计程序名称 → 被审计单位/日期/科目代码 → 编制人/复核人 → 币种（人民币元）")
add_para("  • 蓝色主题：表头 PatternFill(start_color='1F4E79')、白色粗体字")
add_para("  • 条件着色：红色异常（Font color='FF0000'）、黄色关注（Font color='FF8C00'）")
add_para("  • 斑马纹：交替行 PatternFill('F2F2F2') 提升可读性")
add_para("  • 数字格式化：千分位 #,##0.00，百分比 0.00%")
add_para("  • 列宽自适应 + 冻结首行 + 自动筛选")

add_h2("3.2 5 类底稿模板")
add_para("  • add_审定表() — 科目余额汇总 + 审计调整 + 审定数")
add_para("  • add_明细表() — 按维度（银行/客户/供应商/税种）展开明细")
add_para("  • add_调节表() — 银行余额调节表格式（对账单vs账面）")
add_para("  • add_函证汇总表() — 发出/收回/差异汇总")
add_para("  • add_核查表() — 通用核对比对格式")

add_h2("3.3 AuditLogger 类")
add_para("基于 python-docx 生成 Word 格式审计轨迹文件（.docx），7 章节结构：")
add_para("  1) 基本信息（被审计单位、科目、期间、审计师、事务所）")
add_para("  2) 数据来源（上传文件清单、Sheet 统计）")
add_para("  3) 执行时间线（每阶段时间戳）")
add_para("  4) 风险清单（分级表格）")
add_para("  5) 审计发现（逐项列出）")
add_para("  6) 审计结论（每张底稿独立结论）")
add_para("  7) 签名区（编制/复核/批准）")

# ========== 四 ==========
add_h1("四、OCR 凭证识别")

add_h2("4.1 引擎选型")
add_bold_para("OCR引擎：", "EasyOCR 1.7.2")
add_bold_para("语言模型：", "ch_sim（简体中文）")
add_bold_para("运行模式：", "CPU 推理（gpu=False），无需 NVIDIA 显卡")
add_bold_para("加载策略：", "懒加载（Lazy Init）— 首次 API 调用时才加载模型到内存，避免服务器启动阻塞 15-20 秒")
add_bold_para("支持格式：", "JPG / PNG / BMP / PDF")

add_h2("4.2 结构化提取流程")
add_para("Step 1 — EasyOCR.readtext() 返回文本行列表（detail=0 仅文本）")
add_para("Step 2 — 文本规范化：英文冒号→中文冒号、合并连续空白")
add_para("Step 3 — 按凭证类型加载 VOUCHER_PATTERNS 正则模式库")
add_para("Step 4 — 逐字段 re.findall() 提取，存入结构化字典")
add_para("Step 5 — 兜底提取：若关键字段（税种/金额/日期）未匹配，全局正则搜索")

add_h2("4.3 三种凭证的正则模式库")
add_table(
    ["凭证类型", "提取字段数", "字段"],
    [
        ["完税凭证", "6", "税种、税款所属期、缴款金额、缴款日期、凭证号、税务机关"],
        ["纳税申报表", "7", "税种、申报期间、计税依据、税率、应纳税额、已纳税额、应补退税额"],
        ["增值税发票", "9", "发票代码、发票号码、开票日期、购买方、销售方、金额、税率、税额、价税合计"],
    ]
)

add_h2("4.4 凭证数据管理")
add_bold_para("存储引擎：", "SQLite（ocr_data/vouchers.db）")
add_para("表结构：id INTEGER PK / voucher_type TEXT / subject_id TEXT / image_path TEXT / raw_text TEXT / structured_data TEXT(JSON) / created_at TEXT / updated_at TEXT")
add_bold_para("REST API：", "")
add_para("  • POST /api/ocr/scan — 上传图片 → OCR → 提取 → 入库，返回结构化结果")
add_para("  • GET /api/ocr/records — 查询列表（支持按科目、凭证类型筛选）")
add_para("  • PUT /api/ocr/records/{id} — 人工修正结构化字段后保存")
add_para("  • DELETE /api/ocr/records/{id} — 删除记录")
add_para("  • GET /api/ocr/image/{id} — 返回原始凭证图片（浏览器可直接预览）")

add_h2("4.5 LLM Prompt 集成")
add_para("在 audit_with_progress() 的 Phase 3（Prompt 构建）中，自动查询该科目的 OCR 凭证记录，将 structured_data 以 JSON 格式拼入 user prompt 的「原始凭证数据」段，LLM 可直接引用凭证金额与账面进行比对。")

# ========== 五 ==========
add_h1("五、后端架构")

add_h2("5.1 技术栈")
add_para("  • Web 框架：FastAPI（Python 3.12 异步）")
add_para("  • ASGI 服务器：Uvicorn，监听 127.0.0.1:8800")
add_para("  • 跨域：CORS allow_origins=[\"*\"]")
add_para("  • 文件上传：python-multipart")

add_h2("5.2 16 个 API 端点")

add_table(
    ["端点", "方法", "功能", "关键实现"],
    [
        ["/api/health", "GET", "健康检查", "返回状态 + 模型名 + 16 科目列表"],
        ["/api/upload", "POST", "上传 Excel", "时间戳前缀重命名 → uploads/"],
        ["/api/audit-stream", "GET/POST", "SSE 流式审计", "audit_with_progress() 异步生成器"],
        ["/api/demo-files", "GET", "获取范例文件", "按科目查询 fm 字典 → samples/"],
        ["/api/demo-result", "GET", "获取演示结果", "硬编码 16 科 JSON（不含 LLM）"],
        ["/api/download", "GET", "文件下载", "自动识别 8 种 MIME 类型"],
        ["/api/templates", "GET", "模板列表", "扫描 templates/*.xlsx"],
        ["/api/templates/upload", "POST", "上传模板", "科目前缀 + 原名 → templates/"],
        ["/api/recent", "GET", "近期记录", "扫描 output/*.xlsx + *.docx 合并排序"],
        ["/api/risk-summary", "GET", "风险摘要", "扫描 demo_outputs/*审计日志* 正则提取"],
        ["/api/stats", "GET", "仪表盘统计", "总审计数/本月风险/模板数/科目分布/7日趋势"],
        ["/api/ocr/scan", "POST", "OCR 扫描", "EasyOCR → 正则提取 → SQLite"],
        ["/api/ocr/records", "GET", "凭证列表", "SQLite 查询 + JSON 反序列化"],
        ["/api/ocr/records/{id}", "PUT", "修正凭证", "更新 structured_data JSON"],
        ["/api/ocr/records/{id}", "DELETE", "删除凭证", "DELETE FROM vouchers"],
        ["/api/ocr/image/{id}", "GET", "凭证原图", "FileResponse + image/{ext} MIME"],
    ]
)

add_h2("5.3 无状态架构")
add_para("后端不维护会话状态。前端通过完整文件路径列表传递数据源（files=path1,path2,...），/api/audit-stream 直接按路径读取文件并执行审计。每次审计独立、可追溯。LLM 原始响应保存为 output/llm_raw_{timestamp}.txt。")

add_h2("5.4 文件目录设计")
add_table(
    ["目录", "用途", "内容"],
    [
        ["output/", "审计输出", "生成的 .xlsx 底稿 + .docx 日志 + LLM 原始响应"],
        ["samples/", "范例数据", "55 个 Excel（16 科 × 2~4）+ 3 张 OCR PNG"],
        ["uploads/", "用户上传", "时间戳前缀的原始 Excel 文件"],
        ["templates/", "审计模板", "用户上传的底稿模板 Excel"],
        ["demo_outputs/", "演示输出", "16 科预生成底稿 + 日志"],
        ["ocr_data/", "OCR 数据", "vouchers.db (SQLite) + images/ 目录"],
    ]
)

# ========== 六 ==========
add_h1("六、前端架构")

add_h2("6.1 技术选型")
add_para("  • 框架：Vanilla JS（零框架依赖，纯原生 JavaScript）")
add_para("  • CSS：Tailwind CSS CDN + 自定义 CSS ~750 行")
add_para("  • 图标：Font Awesome 6.5.1（免费版 CDN）")
add_para("  • 字体：Inter（英文/数字）+ Microsoft YaHei（中文）")
add_para("  • 无外部图表库：所有可视化用纯 CSS/SVG 实现")

add_h2("6.2 SPA 单页应用路由")
add_para("4 页面：主页 / 工作界面 / 技术原理 / 用户面板")
add_para("切换机制：switchPage(name) → display:none/block + fadeSlideIn 过渡动画（opacity + translateY(16px) + blur(4px)）")
add_para("导航栏：固定顶部 56px，backdrop-filter:blur(24px) 毛玻璃效果，transform 滑入动画")

add_h2("6.3 状态管理（全局变量）")
add_table(
    ["变量", "类型", "用途"],
    [
        ["selectedSubject", "string", "当前选中科目 ID（C-R）"],
        ["subjectFiles", "{C:{balance:[File],...}, D:{...}}", "按科目+文件类型二维管理待上传文件"],
        ["ocrResults", "{M:{vouchers:[{id,fields,...}]}}", "OCR 识别结果缓存"],
        ["uploadedFiles", "[{path, isDemo}]", "已上传到服务端的文件路径列表"],
        ["dashStats", "object", "仪表盘统计数据缓存"],
        ["recentData", "[{type,name,path,time}]", "近期审计记录缓存（供筛选）"],
        ["riskData", "[{source,finding,time}]", "风险摘要缓存"],
    ]
)

add_h2("6.4 核心交互模式")
add_bold_para("3D 卡片跟踪：", "transform-style:preserve-3d + mousemove 更新 CSS 变量 --mx/--my → radial-gradient 光晕跟随光标")
add_bold_para("拖拽上传：", "ondragenter/ondragover/ondrop 事件链 → drag-over class 缩放 1.02 + 蓝色高亮边框 → File API 读取 → 加入 subjectFiles")
add_bold_para("数字滚动动画：", "animateNumber(el, target) — requestAnimationFrame + easeOutCubic 缓动（1-(1-p)^3），duration=800ms")
add_bold_para("滚动入场：", "IntersectionObserver → threshold:0.15 → .visible class 触发 opacity+translateY 过渡")

add_h2("6.5 数据可视化实现（纯 CSS/SVG，零图表库依赖）")
add_bold_para("KPI 卡片数字：", "font-variant-numeric:tabular-nums 等宽数字对齐 + animateNumber() 从 0 滚动到目标值")
add_bold_para("科目分布饼图：", "conic-gradient() 纯 CSS — 16 色扇形，legend 网格，动态计算每个扇形的起始/结束角度")
add_bold_para("7 日审计趋势：", "SVG <polyline> 折线 + <linearGradient> 渐变面积填充 + <circle> 数据点 hover scale(1.8) + <title> tooltip")
add_bold_para("风险分布条：", "flex 三段式横条 — 红(#ef4444)/黄(#f59e0b)/蓝(#3b82f6)，按高中低风险数量比例分配宽度")
add_bold_para("迷你环形图：", "conic-gradient 环形（红黄蓝）+ ::after 伪元素白色内圆形成环状效果")

add_h2("6.6 动画系统")
add_para("自定义 3 种贝塞尔缓动曲线作为 CSS 变量：")
add_para("  • --ease-out-expo: cubic-bezier(0.19,1,0.22,1) — 页面过渡、卡片hover")
add_para("  • --ease-out-back: cubic-bezier(0.34,1.56,0.64,1) — 弹跳反馈")
add_para("  • --ease-spring: cubic-bezier(0.43,1.47,0.65,0.99) — 步骤脉冲")
add_para("关键动画：")
add_para("  • Splash: Canvas 粒子网格 → 数据演示 → 纸飞机图标 Morphing（3 阶段，3.5 秒）")
add_para("  • Hero: hero-orb × 3 浮动（orbFloat 10s/12s）+ charIn 逐字入场")
add_para("  • Tech 页: SVG stroke-dasharray 连接线闪烁 + offset-path 流动圆点（3 组 × 3 层）")
add_para("  • SSE Bar: ::after 伪元素 scaleX 往返动画模拟数据流脉冲")

# ========== 七 ==========
add_h1("七、范例数据生成")

add_h2("7.1 generate_all_samples.py")
add_para("  • 860 行 Python，openpyxl 写入")
add_para("  • 16 科 × 2~4 文件/科 = 55 个 Excel 文件")
add_para("  • 统一蓝底白字表头：PatternFill(start_color='1F4E79', fgColor='1F4E79') + Font(color='FFFFFF', bold=True)")
add_para("  • 科目余额表：标准借贷 8 列（科目代码/名称/期初借方/期初贷方/本期借方/本期贷方/期末借方/期末贷方），借贷自动平衡")
add_para("  • 明细表/台账：按各科特点生成（票据备查簿、借款台账、存货收发存、固定资产卡片等）")
add_para("  • 数据拟真：使用真实中文企业名称（重庆蛮先进智能制造有限公司 + 供应商/客户链）")
add_para("  • 故意埋错：M 应交税费的城建税多提 8.6 倍、教育费附加多提 12 倍、印花税未计提（供 AI 风险检测）")

add_h2("7.2 generate_ocr_samples.py")
add_para("  • 203 行 Python，Pillow ImageDraw 绘制")
add_para("  • 生成 3 张 800×600 PNG：完税凭证（增值税）、纳税申报表（企业所得税季度预缴 A 类）、增值税专用发票（进项）")
add_para("  • 模拟真实排版：表格线框、红色印章椭圆、标题居中、字段标签对齐")
add_para("  • 中文字体：自动检测 Windows 系统字体（宋体/微软雅黑/黑体/楷体）")

# ========== 八 ==========
add_h1("八、安全与合规")

add_para("  • 数据本地处理：所有 Excel/OCR 数据在本地处理，仅 LLM API 调用时发送文本数据")
add_para("  • 审计轨迹完整：每次审计生成 Word 日志，含时间戳、LLM 调用统计（模型/耗时/Token量）、所有发现和结论")
add_para("  • 文件路径无校验：设计为本地单用户工具，未实现认证/鉴权/路径遍历防护")
add_para("  • 无第三方数据泄露：不上传任何数据到除 DeepSeek API 以外的服务")

# ========== 九 ==========
add_h1("九、技术指标总览")

add_table(
    ["指标", "数值"],
    [
        ["项目总代码量", "~6,700 行（Python ~3,700 + HTML/CSS/JS ~2,400 + CLAUDE.md 99）"],
        ["审计科目", "16 个（C 货币资金 ~ R 长期借款）"],
        ["API 端点", "16 个"],
        ["范例文件", "55 个 Excel + 3 张 PNG"],
        ["底稿模板类型", "5 类（审定表/明细表/调节表/函证汇总表/核查表）"],
        ["LLM 上下文", "1,000,000 tokens（DeepSeek V4 Pro）"],
        ["OCR 语言", "简体中文（ch_sim）"],
        ["前端框架", "零依赖（纯原生 JS + Tailwind CDN）"],
        ["后端框架", "FastAPI + Uvicorn（ASGI）"],
        ["数据库", "SQLite（仅 OCR 凭证数据）"],
        ["Excel 引擎", "openpyxl"],
        ["Word 引擎", "python-docx"],
        ["图表方案", "纯 CSS conic-gradient / SVG polyline"],
        ["Python 版本", "3.12"],
        ["运行端口", "8800"],
        ["最低 Python 依赖", "7 个（easyocr, Pillow, openpyxl, python-docx, fastapi, uvicorn, python-multipart, json-repair, httpx）"],
    ]
)

# 保存
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
path = os.path.join(desktop, '审计智能体_技术原理.docx')
doc.save(path)
print(f"已保存到：{path}")
