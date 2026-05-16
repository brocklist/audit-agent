"""
审计智能体 - FastAPI 后端 v4
LLM分析数据 → 服务端模板渲染Excel → 输出审计日志
"""
import json, os, re, io, asyncio, time
from datetime import datetime
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, numbers

# ==================== 审计日志系统 ====================
class AuditLogger:
    """收集审计全流程日志，输出可追溯的审计轨迹文件"""
    def __init__(self, entity, period, subject, subject_name, auditor, firm):
        self.entity = entity; self.period = period
        self.subject = subject; self.subject_name = subject_name
        self.auditor = auditor; self.firm = firm
        self.start_time = time.time()
        self.events = []  # [{timestamp, phase, type, message, detail}]
        self.data_summary = {}  # 文件→{sheets, rows, cols}
        self.risk_items = []  # 风险数据清单
        self.audit_findings = []
        self.conclusions = []
        self.llm_stats = {}  # prompt_size, response_size, model, elapsed
        self.output_excel = ""; self.output_log = ""

    def log(self, phase, type, message, detail=""):
        self.events.append({
            "ts": datetime.now().strftime("%H:%M:%S"),
            "elapsed": round(time.time() - self.start_time, 1),
            "phase": phase, "type": type,
            "message": message, "detail": detail
        })

    def add_file(self, name, sheets, rows, preview=""):
        self.data_summary[name] = {"sheets": sheets, "rows": rows, "preview": preview}

    def add_risk(self, level, item, amount, reason):
        self.risk_items.append({"level": level, "item": item, "amount": amount, "reason": reason})

    def add_finding(self, finding):
        self.audit_findings.append(finding)

    def add_conclusion(self, sheet_name, text):
        self.conclusions.append({"sheet": sheet_name, "conclusion": text})

    def set_llm_stats(self, model, prompt_size, response_size, elapsed):
        self.llm_stats = {"model": model, "prompt_size": prompt_size,
                          "response_size": response_size, "elapsed": elapsed}

    def generate_log_file(self) -> str:
        """生成 Word(.docx) 格式的审计日志文件"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)

        elapsed_total = round(time.time() - self.start_time, 1)

        def add_heading(text, level=1):
            h = doc.add_heading(text, level=level)
            for run in h.runs: run.font.name = '微软雅黑'

        def add_table(headers, rows):
            table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Grid Accent 1')
            for i, h in enumerate(headers): table.rows[0].cells[i].text = str(h)
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row): table.rows[ri+1].cells[ci].text = str(val)
            doc.add_paragraph()

        add_heading('审计日志 - AI Audit Trail', 0)

        add_heading('审计基本信息', 1)
        add_table(['项目','内容'], [['被审计单位',self.entity],['审计科目',f'{self.subject_name}（索引号:{self.subject}）'],['会计期间',self.period],['审核员',self.auditor],['事务所',self.firm],['审计日期',datetime.now().strftime('%Y-%m-%d')],['总耗时',f'{elapsed_total:.1f}秒']])

        add_heading('一、数据源读取', 1)
        if self.data_summary:
            add_table(['文件名','Sheet数','总行数','内容摘要'], [[fn,str(i['sheets']),str(i['rows']),i.get('preview','')[:80]] for fn,i in self.data_summary.items()])

        add_heading('二、审计执行时间线', 1)
        add_table(['时间','耗时(s)','阶段','类型','操作','详情'], [[e['ts'],str(e['elapsed']),e['phase'],e['type'],e['message'][:50],e['detail'][:80]] for e in self.events])

        if self.llm_stats:
            add_heading('三、AI模型调用统计', 1)
            add_table(['指标','数值'], [['模型',self.llm_stats.get('model','')],['Prompt大小',f"{self.llm_stats.get('prompt_size',0):,}字符"],['响应大小',f"{self.llm_stats.get('response_size',0):,}字符"],['AI耗时',f"{self.llm_stats.get('elapsed',0):.1f}秒"]])

        add_heading('四、风险数据清单', 1)
        if self.risk_items:
            add_table(['风险等级','科目/项目','涉及金额','风险原因'], [['高' if r['level']=='error' else '中' if r['level']=='warning' else '低',r['item'],str(r['amount']),r['reason']] for r in self.risk_items])
        else:
            doc.add_paragraph('未发现重大风险事项')

        add_heading('五、审计发现汇总', 1)
        for i, f in enumerate(self.audit_findings, 1): doc.add_paragraph(f'{i}. {f}')

        add_heading('六、各工作底稿审计结论', 1)
        for c in self.conclusions:
            add_heading(c['sheet'], 2)
            doc.add_paragraph(c['conclusion'])

        add_heading('七、审计签名', 1)
        add_table(['角色','签名','日期'], [['编制人（AI审计智能体）','',datetime.now().strftime('%Y-%m-%d')],['复核人','',''],['项目负责人','','']])

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r'[\\/*?:"<>|]', '_', self.entity or "审计")
        fname = f"{safe}_{self.period.replace('-','')}_审计日志_{ts}.docx"
        path = OUTPUT_DIR / fname
        doc.save(str(path))
        self.output_log = str(path)
        return str(path)
from openpyxl.utils import get_column_letter
from fastapi import FastAPI, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
import httpx
from json_repair import repair_json

app = FastAPI(title="审计智能体")
BASE_DIR = Path(__file__).parent
CLAUDE_MD = BASE_DIR / "CLAUDE.md"
OUTPUT_DIR = BASE_DIR / "output"; UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR.mkdir(exist_ok=True); UPLOAD_DIR.mkdir(exist_ok=True)

ANTHROPIC_BASE_URL = os.getenv("ANTHROPIC_BASE_URL", "https://api.deepseek.com/anthropic")
ANTHROPIC_AUTH_TOKEN = os.getenv("ANTHROPIC_AUTH_TOKEN", "")
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "deepseek-v4-pro[1m]")

SUBJECT_MAP = {
    "C": {"name": "货币资金", "codes": ["1001","1002","1003","1012"]},
    "D": {"name": "应收票据", "codes": ["1121"]},
    "E": {"name": "其他应收款", "codes": ["1221"]},
    "F": {"name": "预付账款", "codes": ["1123"]},
    "G": {"name": "预收账款", "codes": ["2203"]},
    "H": {"name": "短期借款", "codes": ["2001"]},
}

class AuditRequest(BaseModel):
    subject: str = "C"; entity_name: str = ""; period: str = "2024-12-31"
    auditor: str = ""; reviewer: str = ""; firm_name: str = ""; files: list[str] = []

# ==================== 样式常量 ====================
THIN = Side(style='thin', color='BFBFBF'); MEDIUM = Side(style='medium', color='1F497D')
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
BORDER_HEADER = Border(left=MEDIUM, right=MEDIUM, top=THIN, bottom=MEDIUM)

FILL_TITLE = PatternFill("solid", fgColor="1F497D")
FILL_SUB = PatternFill("solid", fgColor="2E75B6")
FILL_LIGHT = PatternFill("solid", fgColor="D6E4F0")
FILL_GREEN = PatternFill("solid", fgColor="E2EFDA")
FILL_RED = PatternFill("solid", fgColor="FCE4D6")
FILL_YELLOW = PatternFill("solid", fgColor="FFFBE6")
FILL_GREY = PatternFill("solid", fgColor="F2F2F2")
FILL_WHITE = PatternFill("solid", fgColor="FFFFFF")
FILL_INFO = PatternFill("solid", fgColor="EBF3FB")

FT_TITLE = Font(name="微软雅黑", bold=True, size=14, color="FFFFFF")
FT_SEC = Font(name="微软雅黑", bold=True, size=11, color="FFFFFF")
FT_HEADER = Font(name="微软雅黑", bold=True, size=10, color="FFFFFF")
FT_BOLD = Font(name="微软雅黑", bold=True, size=10)
FT_NORM = Font(name="微软雅黑", size=10)
FT_SMALL = Font(name="微软雅黑", size=9)
FT_SMALL_GREY = Font(name="微软雅黑", size=9, color="666666")
FT_GREEN = Font(name="微软雅黑", bold=True, size=10, color="375623")
FT_RED = Font(name="微软雅黑", bold=True, size=10, color="C00000")
FT_BLUE_BOLD = Font(name="微软雅黑", bold=True, size=10, color="1F497D")

AL_C = Alignment(horizontal="center", vertical="center", wrap_text=True)
AL_L = Alignment(horizontal="left", vertical="center", wrap_text=True)
AL_R = Alignment(horizontal="right", vertical="center")
AL_W = Alignment(horizontal="left", vertical="center", wrap_text=True)

# ==================== Excel读取 ====================
def read_excel(path: str) -> dict:
    wb = openpyxl.load_workbook(path, data_only=True)
    result = {}
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = []
        for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 500),
                                max_col=min(ws.max_column, 30), values_only=True):
            rows.append([str(c) if c is not None else "" for c in row])
        result[sn] = {"max_row": ws.max_row, "max_col": ws.max_column, "rows": rows}
    wb.close()
    return result

def serialize_for_llm(data: dict) -> str:
    parts = []
    for sn, sd in data.items():
        parts.append(f"\n=== Sheet: {sn} ({sd['max_row']}r x {sd['max_col']}c) ===")
        for i, row in enumerate(sd["rows"]):
            parts.append(f"  R{i+1}: {row}")
    return "\n".join(parts)

# ==================== LLM调用 ====================
async def call_llm(system_prompt: str, user_message: str) -> str:
    headers = {"Content-Type": "application/json"}
    if ANTHROPIC_AUTH_TOKEN:
        headers["Authorization"] = f"Bearer {ANTHROPIC_AUTH_TOKEN}"
    body = {"model": ANTHROPIC_MODEL, "max_tokens": 16000, "temperature": 0.05,
            "thinking": {"type": "disabled"}, "system": system_prompt,
            "messages": [{"role": "user", "content": user_message}]}
    async with httpx.AsyncClient(timeout=600.0) as client:
        resp = await client.post(f"{ANTHROPIC_BASE_URL}/messages", headers=headers, json=body)
        if resp.status_code != 200:
            raise HTTPException(500, f"LLM API error: {resp.status_code} {resp.text[:300]}")
        data = resp.json()
        for block in data["content"]:
            if block.get("type") == "text": return block["text"]
        return data["content"][0].get("text", "")

def extract_json(text: str) -> dict:
    m = re.search(r'```json\s*([\s\S]*?)\s*```', text)
    raw = m.group(1) if m else (re.search(r'\{[\s\S]*\}', text).group(0) if re.search(r'\{[\s\S]*\}', text) else "{}")
    try: return json.loads(raw)
    except json.JSONDecodeError: return json.loads(repair_json(raw))

# ==================== Excel模板渲染 ====================
def safe_sheet_name(s, default="Sheet"):
    return re.sub(r'[\\/*?:\[\]]', '-', s or default)[:31]

class ExcelBuilder:
    """生成排版规范的审计底稿Excel"""
    def __init__(self, entity, period, auditor, reviewer, firm, index):
        self.wb = openpyxl.Workbook()
        self.wb.remove(self.wb.active)
        self.entity = entity; self.period = period; self.auditor = auditor
        self.reviewer = reviewer; self.firm = firm; self.index = index
        self.date_str = datetime.now().strftime("%Y-%m-%d")

    def _sheet_header(self, ws, title, ncols, sheet_no=""):
        """统一的Sheet页眉：标题行 + 审计信息行"""
        idx = f"{self.index}{f'-{sheet_no}' if sheet_no else ''}"
        # Row 1: 标题
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
        c = ws.cell(1, 1, title); c.font = FT_TITLE; c.fill = FILL_TITLE; c.alignment = AL_C
        ws.row_dimensions[1].height = 36
        # Row 2: 事务所 + 索引号
        ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=ncols//2)
        ws.cell(2, 1, f"事务所：{self.firm or 'XX会计师事务所'}").font = FT_SMALL
        ws.cell(2, 1).fill = FILL_LIGHT; ws.cell(2, 1).alignment = AL_L
        ws.merge_cells(start_row=2, start_column=ncols//2+1, end_row=2, end_column=ncols)
        ws.cell(2, ncols//2+1, f"索引号：{idx}").font = FT_SMALL
        ws.cell(2, ncols//2+1).fill = FILL_LIGHT; ws.cell(2, ncols//2+1).alignment = AL_R
        ws.row_dimensions[2].height = 20
        # Row 3: 被审计单位 + 审核员
        ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=ncols//2)
        ws.cell(3, 1, f"被审计单位：{self.entity}").font = FT_BOLD
        ws.cell(3, 1).fill = FILL_INFO; ws.cell(3, 1).alignment = AL_L
        ws.merge_cells(start_row=3, start_column=ncols//2+1, end_row=3, end_column=ncols)
        ws.cell(3, ncols//2+1, f"审核员：{self.auditor}").font = FT_SMALL
        ws.cell(3, ncols//2+1).fill = FILL_INFO; ws.cell(3, ncols//2+1).alignment = AL_R
        ws.row_dimensions[3].height = 20
        # Row 4: 会计期间 + 日期
        ws.merge_cells(start_row=4, start_column=1, end_row=4, end_column=ncols//2)
        ws.cell(4, 1, f"会计期间：{self.period}").font = FT_BOLD
        ws.cell(4, 1).fill = FILL_INFO; ws.cell(4, 1).alignment = AL_L
        ws.merge_cells(start_row=4, start_column=ncols//2+1, end_row=4, end_column=ncols)
        ws.cell(4, ncols//2+1, f"日期：{self.date_str}").font = FT_SMALL
        ws.cell(4, ncols//2+1).fill = FILL_INFO; ws.cell(4, ncols//2+1).alignment = AL_R
        ws.row_dimensions[4].height = 20

    def _data_header(self, ws, row, headers):
        """写表头行"""
        ws.row_dimensions[row].height = 24
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row, c, h)
            cell.font = FT_HEADER; cell.fill = FILL_SUB; cell.alignment = AL_C; cell.border = BORDER

    def _data_row(self, ws, row, values, formats=None, fills=None, fonts=None):
        """写一行数据"""
        ws.row_dimensions[row].height = 20
        for c, v in enumerate(values, 1):
            cell = ws.cell(row, c, v)
            cell.font = fonts[c-1] if fonts and c-1 < len(fonts) else FT_NORM
            cell.fill = fills[c-1] if fills and c-1 < len(fills) else FILL_WHITE
            cell.alignment = AL_C if c <= 1 else AL_R if isinstance(v, (int, float)) else AL_L
            cell.border = BORDER
            if formats and c-1 < len(formats) and formats[c-1]:
                cell.number_format = formats[c-1]

    def _conclusion(self, ws, row, text, ncols):
        """写审计结论行"""
        ws.merge_cells(start_row=row, start_column=1, end_row=row+2, end_column=ncols)
        c = ws.cell(row, 1, f"审计结论：{text}")
        c.font = FT_SMALL_GREY; c.fill = FILL_YELLOW; c.alignment = AL_W; c.border = BORDER
        ws.row_dimensions[row].height = 45; ws.row_dimensions[row+1].height = 25

    # ---- 具体Sheet模板 ----

    def add_审定表(self, data: dict):
        """货币资金/应收票据等审定表"""
        headers = data.get("headers", ["序号","科目代码","科目名称","期初余额","期末余额","变动额","变动率","审计调整","审定数"])
        ncols = len(headers)
        title = safe_sheet_name(data.get("title"), "审定表")
        ws = self.wb.create_sheet(title)
        self._sheet_header(ws, data.get("title", "审定表"), ncols)

        # 列宽
        widths = data.get("col_widths", [4,10,26,15,15,15,12,14,15])
        for i, w in enumerate(widths[:ncols], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # 表头
        self._data_header(ws, 5, headers)

        # 数据行
        rows = data.get("rows", [])
        num_fmt = '#,##0.00'
        pct_fmt = '0.00%'
        for i, row in enumerate(rows):
            r = 6 + i
            fills = [FILL_WHITE if i%2==0 else FILL_GREY] * ncols
            # 异常变动率标红/黄
            rate_idx = headers.index("变动率") if "变动率" in headers else -1
            if rate_idx >= 0 and rate_idx < len(row):
                try:
                    rate_str = str(row[rate_idx]).replace('%','')
                    rate = float(rate_str) / 100 if '.' in rate_str else float(rate_str)
                    if abs(rate) > 50: fills = [FILL_RED] * ncols
                    elif abs(rate) > 30: fills = [FILL_YELLOW] * ncols
                except: pass

            formats = [None] * ncols
            for j, h in enumerate(headers):
                if h in ("期初余额","期末余额","变动额","审计调整","审定数","金额","借方金额","贷方金额","票面金额","坏账准备","已计提坏账","应计提坏账"):
                    formats[j] = num_fmt
                elif h in ("变动率","结构","占比","利率","比例"):
                    formats[j] = pct_fmt

            self._data_row(ws, r, row, formats=formats, fills=fills)

        # 结论
        next_row = 6 + len(rows) + 1
        self._conclusion(ws, next_row, data.get("conclusion", ""), ncols)
        return ws

    def add_明细表(self, data: dict):
        """明细表：应收票据/其他应收款/预付账款/预收账款/短期借款"""
        headers = data.get("headers", [])
        ncols = len(headers)
        if ncols == 0: return None

        ws = self.wb.create_sheet(safe_sheet_name(data.get('title'),'明细表'))
        self._sheet_header(ws, data.get("title", "明细表"), ncols)

        widths = data.get("col_widths", [4]*ncols)
        for i, w in enumerate(widths[:ncols], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        self._data_header(ws, 5, headers)
        rows = data.get("rows", [])
        for i, row in enumerate(rows):
            r = 6 + i
            fills = [FILL_WHITE if i%2==0 else FILL_GREY] * ncols
            # 逾期/异常标红
            for j, v in enumerate(row):
                sv = str(v)
                if "逾期" in sv or "长期" in sv or "挂账" in sv or "未回函" in sv or "差异" in sv:
                    fills = [FILL_RED if k==j else f for k,f in enumerate(fills)]
            self._data_row(ws, r, row, fills=fills)

        next_row = 6 + len(rows) + 1
        self._conclusion(ws, next_row, data.get("conclusion", ""), ncols)
        return ws

    def add_调节表(self, data: dict):
        """银行存款余额调节表"""
        ws = self.wb.create_sheet(safe_sheet_name(data.get('title'),'余额调节表'))
        ncols = 6
        self._sheet_header(ws, data.get("title", "余额调节表"), ncols)
        for c, w in enumerate([4,30,16,16,16,30], 1):
            ws.column_dimensions[get_column_letter(c)].width = w

        # 银行对账单部分
        self._data_header(ws, 5, ["序号","项目","金额","","序号","项目","金额"])
        bank_items = data.get("bank_items", [])
        total_bank = 0
        for i, (label, amt) in enumerate(bank_items):
            r = 6 + i
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
            ws.cell(r, 1, label).font = FT_BOLD; ws.cell(r, 1).alignment = AL_L
            ws.cell(r, 3, amt).font = FT_NORM; ws.cell(r, 3).number_format = '#,##0.00'; ws.cell(r, 3).alignment = AL_R
            for cc in range(1, 7): ws.cell(r, cc).border = BORDER; ws.cell(r, cc).fill = FILL_WHITE if i%2==0 else FILL_GREY

        # 企业部分
        enterprise_items = data.get("enterprise_items", [])
        for i, (label, amt) in enumerate(enterprise_items):
            r = 6 + i
            ws.cell(r, 4, label).font = FT_BOLD; ws.cell(r, 4).alignment = AL_L
            ws.cell(r, 6, amt).font = FT_NORM; ws.cell(r, 6).number_format = '#,##0.00'; ws.cell(r, 6).alignment = AL_R
            for cc in range(1, 7): ws.cell(r, cc).border = BORDER; ws.cell(r, cc).fill = FILL_WHITE if i%2==0 else FILL_GREY

        conclusion = data.get("conclusion", "")
        r = 6 + max(len(bank_items), len(enterprise_items)) + 1
        self._conclusion(ws, r, conclusion, 6)
        return ws

    def add_函证汇总表(self, data: dict):
        """银行函证/票据函证汇总"""
        headers = data.get("headers", ["序号","银行/出票人","账号/票据号","币种","账面余额","函证金额","是否回函","差异"])
        ncols = len(headers)
        ws = self.wb.create_sheet(safe_sheet_name(data.get('title'), '函证汇总表'))
        self._sheet_header(ws, data.get("title", "函证汇总表"), ncols)
        widths = data.get("col_widths", [4,22,16,8,15,15,10,18])
        for i, w in enumerate(widths[:ncols], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        self._data_header(ws, 5, headers)
        rows = data.get("rows", [])
        for i, row in enumerate(rows):
            r = 6 + i
            fills = [FILL_WHITE if i%2==0 else FILL_GREY] * ncols
            # 未回函标黄，有差异标红
            for j, v in enumerate(row):
                sv = str(v)
                if "未回函" in sv: fills = [FILL_YELLOW if k==j else f for k,f in enumerate(fills)]
                if "差异" in sv or "不符" in sv: fills = [FILL_RED if k==j else f for k,f in enumerate(fills)]
            self._data_row(ws, r, row, fills=fills)

        next_row = 6 + len(rows) + 1
        self._conclusion(ws, next_row, data.get("conclusion", ""), ncols)
        return ws

    def add_核查表(self, data: dict):
        """大额流水/收入确认等核查表"""
        headers = data.get("headers", [])
        ncols = len(headers)
        if ncols == 0: return None
        ws = self.wb.create_sheet(safe_sheet_name(data.get('title'), '核查表'))
        self._sheet_header(ws, data.get("title", "核查表"), ncols)
        widths = data.get("col_widths", [4]*ncols)
        for i, w in enumerate(widths[:ncols], 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        self._data_header(ws, 5, headers)
        rows = data.get("rows", [])
        for i, row in enumerate(rows):
            r = 6 + i
            fills = [FILL_WHITE if i%2==0 else FILL_GREY] * ncols
            for j, v in enumerate(row):
                sv = str(v)
                if "大额" in sv or "异常" in sv or "风险" in sv or "需核实" in sv:
                    fills = [FILL_YELLOW if k==j else f for k,f in enumerate(fills)]
                if "关注" in sv or "逾期" in sv:
                    fills = [FILL_RED if k==j else f for k,f in enumerate(fills)]
            self._data_row(ws, r, row, fills=fills)

        next_row = 6 + len(rows) + 1
        self._conclusion(ws, next_row, data.get("conclusion", ""), ncols)
        return ws

    def save(self):
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = re.sub(r'[\\/*?:"<>|]', '_', self.entity or "审计底稿")
        fname = f"{safe}_{self.period.replace('-','')}_{ts}.xlsx"
        path = OUTPUT_DIR / fname
        self.wb.save(str(path)); self.wb.close()
        return str(path)


def build_excel_from_data(audit_data: dict, entity: str, period: str,
                          auditor: str, reviewer: str, firm: str, index: str) -> str:
    """从结构化审计数据构建Excel底稿"""
    builder = ExcelBuilder(entity, period, auditor, reviewer, firm, index)
    sheets = audit_data.get("sheets", [])

    for sdef in sheets:
        stype = sdef.get("type", "审定表")
        # 所有表类型统一使用审定表渲染（header+rows+conclusion模式）
        # 特殊类型保留路由
        if stype in ("调节表",):
            builder.add_调节表(sdef)
        elif stype in ("函证", "函证汇总表"):
            builder.add_函证汇总表(sdef)
        else:
            # 审定表/明细表/账龄分析表/盘点表/利息测算/截止测试/核查表 等均使用统一模板
            builder.add_审定表(sdef)

    return builder.save()


# ==================== SSE进度 ====================
async def audit_with_progress(req: AuditRequest):
    start_time = time.time()
    subj = SUBJECT_MAP.get(req.subject, SUBJECT_MAP["C"])

    # 创建日志器
    logger = AuditLogger(req.entity_name or "被审计单位", req.period,
                         req.subject, subj['name'],
                         req.auditor or "AI审计智能体", req.firm_name or "XX会计师事务所")

    def emit(pct, phase, step, detail, status="active"):
        e = time.time() - start_time
        logger.log(phase, status, step, detail)
        return f"event: progress\ndata: " + json.dumps({
            "pct": pct, "phase": phase, "step": step, "detail": detail,
            "status": status, "elapsed": round(e, 1)
        }, ensure_ascii=False) + "\n\n"

    # Phase 1: 读文件
    logger.log("file_read", "info", "开始读取数据源", f"共 {len(req.files)} 个文件")
    yield emit(2, "file_read", "扫描数据源", f"发现 {len(req.files)} 个文件")
    await asyncio.sleep(0.1)
    all_data = {}; total_rows = 0
    for i, fp in enumerate(req.files):
        if not os.path.exists(fp):
            logger.log("file_read", "warn", f"文件不存在: {os.path.basename(fp)}", "")
            yield emit(3+i, "file_read", os.path.basename(fp), "文件不存在，跳过", "warn")
            continue
        try:
            data = read_excel(fp)
            all_data[os.path.basename(fp)] = data
            rows = sum(d['max_row'] for d in data.values())
            total_rows += rows
            # 记录文件摘要
            first_rows = []
            for sd in data.values():
                for r in sd['rows'][:3]:
                    first_rows.append(str(r)[:100])
                    break
                break
            logger.add_file(os.path.basename(fp), len(data), rows,
                           str(first_rows[0])[:80] if first_rows else "")
            yield emit(4+i*2, "file_read", f"加载: {os.path.basename(fp)[:30]}", f"{len(data)} Sheet · {rows} 行", "done")
        except Exception as e:
            logger.log("file_read", "error", f"读取失败: {os.path.basename(fp)}", str(e)[:60])
            yield emit(4+i, "file_read", os.path.basename(fp), str(e)[:40], "error")
    if not all_data:
        yield "event: error\ndata: " + json.dumps({"msg": "未能读取有效数据"}, ensure_ascii=False) + "\n\n"; return
    logger.log("file_read", "done", "数据读取完成", f"{len(all_data)} 文件 · {total_rows} 行")
    yield emit(12, "file_read", "数据读取完成", f"{len(all_data)} 文件 · {total_rows} 行", "done")
    await asyncio.sleep(0.1)

    # Phase 2: 数据序列化
    logger.log("data_prep", "info", "数据序列化", "将Excel转为AI可读文本")
    yield emit(13, "data_prep", "数据序列化", "转换为AI可读文本...")
    await asyncio.sleep(0.1)
    data_text = ""
    for fn, fd in all_data.items():
        data_text += f"\n{'='*60}\n文件: {fn}\n{'='*60}\n" + serialize_for_llm(fd)
    logger.log("data_prep", "done", "序列化完成", f"{len(data_text):,} 字符")
    yield emit(16, "data_prep", "数据序列化完成", f"{len(data_text):,} 字符")
    await asyncio.sleep(0.1)

    rules = CLAUDE_MD.read_text(encoding="utf-8") if CLAUDE_MD.exists() else ""
    logger.log("data_prep", "done", "加载审计准则", f"{subj['name']} · {len(rules):,} 字符")
    yield emit(18, "data_prep", "加载审计准则", f"{subj['name']} · 代码 {','.join(subj['codes'])}")
    await asyncio.sleep(0.1)

    # Phase 3: 构建Prompt
    logger.log("prompt", "info", "构建审计指令", "组装系统提示+用户数据")
    yield emit(20, "prompt", "构建审计指令", "组装系统提示与用户数据...")
    await asyncio.sleep(0.1)

    user_prompt = f"""## 审计任务
**被审计单位**: {req.entity_name or '未指定'}
**审计科目**: {subj['name']}（科目代码: {','.join(subj['codes'])}）
**会计期间**: {req.period}
**审核员**: {req.auditor or 'AI审计智能体'}
**事务所**: {req.firm_name or 'XX会计师事务所'}

## 原始数据
{data_text}

## 输出要求
请严格按照 CLAUDE.md 中的审计程序执行审计，返回以下结构的JSON（**不要返回cell级别的Excel布局，只返回数据**）:
```json
{{
  "summary": "整体审计发现（200字内）",
  "risk_items": [
    {{"level": "error或warning或info", "item": "科目/项目名称", "amount": "涉及金额", "reason": "风险原因"}}
  ],
  "findings": ["审计发现1", "审计发现2", ...],
  "sheets": [
    {{
      "type": "审定表",
      "title": "C-1 货币资金审定表",
      "headers": ["序号","科目代码","科目名称","期初余额","期末余额","变动额","变动率","审计调整","审定数"],
      "col_widths": [4,10,26,15,15,15,12,14,15],
      "rows": [["1","1001","库存现金",50000.00,55000.00,5000.00,"10.00%",0.00,55000.00], ...],
      "conclusion": "经审计，..."
    }}
  ]
}}
```

**关键要求**:
1. 除了 summary 和 sheets 外，还必须返回 risk_items 和 findings 字段
2. risk_items 列出所有异常/风险项目，level取error(高风险)/warning(中风险)/info(低风险)
3. sheets 中每个sheet.type取: "审定表"/"明细"/"调节表"/"函证"/"核查"
4. rows是二维数组，每行与headers一一对应，金额两位小数，百分比字符串格式
5. 变动率>30%时标warning，>50%标error
6. 每个sheet必须有conclusion"""

    prompt_size = len(user_prompt)
    logger.log("prompt", "done", "Prompt组装完成", f"系统 {len(rules):,} + 用户 {prompt_size:,} 字符")
    yield emit(23, "prompt", "Prompt组装完成", f"系统 {len(rules):,} + 用户 {prompt_size:,} 字符")
    await asyncio.sleep(0.2)

    # Phase 4: AI分析
    logger.log("ai_analysis", "info", "发起AI审计请求", f"模型: {ANTHROPIC_MODEL}")
    yield emit(25, "ai_analysis", "发起AI审计请求", f"模型: {ANTHROPIC_MODEL}")
    llm_start = time.time()
    llm_task = asyncio.create_task(call_llm(rules, user_prompt))

    stages = [(8,28,"科目余额分析"), (12,33,"明细账审查"), (15,40,"对账单核对"),
              (18,48,"函证检查"), (20,55,"编制审定表"), (22,62,"编制明细表"),
              (25,70,"编写审计发现"), (28,78,"编写审计结论"), (30,84,"最终整理")]
    for delay, pct, step in stages:
        await asyncio.sleep(delay)
        if llm_task.done(): break
        logger.log("ai_analysis", "active", step, "AI深度分析中")
        yield emit(pct, "ai_analysis", step, "AI深度分析中...")

    extra = 0
    while not llm_task.done() and extra < 360:
        await asyncio.sleep(20); extra += 20
        if llm_task.done(): break
        logger.log("ai_analysis", "active", "AI深度分析中", f"已等待 {int(time.time()-llm_start)} 秒")
        yield emit(min(85, 80+extra//20), "ai_analysis", "AI深度分析中...", f"已等待 {int(time.time()-start_time)} 秒")

    try:
        llm_response = await llm_task
        llm_elapsed = time.time() - llm_start
        response_size = len(llm_response)
        logger.set_llm_stats(ANTHROPIC_MODEL, prompt_size, response_size, llm_elapsed)
        logger.log("ai_analysis", "done", "AI分析完成", f"耗时 {llm_elapsed:.0f}s · 响应 {response_size:,} 字符")
        yield emit(88, "ai_analysis", "AI分析完成", f"耗时 {llm_elapsed:.0f}s · 响应 {response_size:,} 字符", "done")
    except Exception as e:
        logger.log("ai_analysis", "error", "AI调用失败", str(e)[:60])
        yield "event: error\ndata: " + json.dumps({"msg": f"AI调用失败: {e}"}, ensure_ascii=False) + "\n\n"; return

    # Phase 5: 解析结果
    logger.log("parsing", "info", "解析审计结果", "提取结构化数据...")
    yield emit(89, "parsing", "解析审计结果", "提取结构化数据...")
    await asyncio.sleep(0.1)
    try:
        audit_result = extract_json(llm_response)
    except ValueError as e:
        logger.log("parsing", "error", "JSON解析失败", str(e)[:80])
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        (OUTPUT_DIR / f"llm_raw_{ts}.txt").write_text(llm_response, encoding="utf-8")
        yield "event: error\ndata: " + json.dumps({"msg": f"JSON解析失败: {e}"}, ensure_ascii=False) + "\n\n"; return

    # 记录风险项
    for r in audit_result.get("risk_items", []):
        logger.add_risk(r.get("level","info"), r.get("item",""), r.get("amount",""), r.get("reason",""))
    # 记录审计发现
    for f in audit_result.get("findings", []):
        logger.add_finding(f)

    sheet_count = len(audit_result.get("sheets", []))
    logger.log("parsing", "done", "数据结构化完成", f"{sheet_count} 张Sheet待生成 · {len(audit_result.get('risk_items',[]))} 条风险")
    yield emit(92, "parsing", "数据结构化完成", f"{sheet_count} 张Sheet · {len(audit_result.get('risk_items',[]))} 条风险", "done")
    await asyncio.sleep(0.1)

    # Phase 6: 渲染Excel
    for i, sdef in enumerate(audit_result.get("sheets", [])):
        pct = 93 + i * 6 // max(sheet_count, 1)
        title = sdef.get("title", "")[:35]
        logger.log("generating", "info", f"渲染Sheet: {title}", f"第 {i+1}/{sheet_count} 张 · type={sdef.get('type','')}")
        logger.add_conclusion(sdef.get("title",""), sdef.get("conclusion",""))
        yield emit(pct, "generating", f"渲染: {title}", f"第 {i+1}/{sheet_count} 张工作底稿")
        await asyncio.sleep(0.03)

    try:
        out_path = build_excel_from_data(audit_result, req.entity_name or "被审计单位",
                                         req.period, req.auditor or "", req.reviewer or "",
                                         req.firm_name or "", req.subject)
        logger.output_excel = out_path
    except Exception as e:
        logger.log("generating", "error", "Excel生成失败", str(e)[:80])
        yield "event: error\ndata: " + json.dumps({"msg": f"Excel生成失败: {e}"}, ensure_ascii=False) + "\n\n"; return

    # 生成审计日志文件
    log_path = logger.generate_log_file()
    logger.log("complete", "done", "审计日志已生成", log_path)
    yield emit(99, "complete", "生成审计日志", f"日志文件: {os.path.basename(log_path)}")

    elapsed = time.time() - start_time
    logger.log("complete", "done", "审计任务完成",
               f"{sheet_count} 张底稿 · {len(logger.risk_items)} 条风险 · {elapsed:.0f}秒")

    yield emit(100, "complete", "审计任务完成",
               f"{sheet_count} 张底稿 · {len(logger.risk_items)} 条风险 · {elapsed:.0f}秒", "done")

    yield "event: result\ndata: " + json.dumps({
        "success": True,
        "output_file": out_path,
        "log_file": log_path,
        "summary": audit_result.get("summary", ""),
        "sheets": [s["title"] for s in audit_result.get("sheets", [])],
        "risk_count": len(logger.risk_items),
        "log_content": open(log_path, encoding='utf-8').read()[:3000],
        "elapsed_total": round(elapsed, 1)
    }, ensure_ascii=False) + "\n\n"


# ==================== API路由 ====================
@app.post("/api/upload")
async def upload_files(files: list[UploadFile]):
    saved = []
    for f in files:
        if not f.filename.endswith(('.xlsx', '.xls')): raise HTTPException(400, f"仅Excel: {f.filename}")
        content = await f.read()
        ts = datetime.now().strftime("%H%M%S")
        path = UPLOAD_DIR / f"{ts}_{re.sub(r'[\\/*?:<>|]','_',f.filename)}"
        path.write_bytes(content); saved.append(str(path))
    return {"success": True, "files": saved, "count": len(saved)}

@app.get("/api/audit-stream")
async def audit_stream(entity_name: str="", period: str="2024-12-31", auditor: str="",
                       reviewer: str="", firm_name: str="", subject: str="C", files: str=""):
    fl = [f.strip() for f in files.split(",") if f.strip()]
    req = AuditRequest(subject=subject, entity_name=entity_name, period=period,
                       auditor=auditor, reviewer=reviewer, firm_name=firm_name, files=fl)
    return StreamingResponse(audit_with_progress(req), media_type="text/event-stream",
                            headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

@app.post("/api/audit-stream")
async def audit_stream_post(req: AuditRequest):
    return StreamingResponse(audit_with_progress(req), media_type="text/event-stream",
                            headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

@app.get("/api/demo-result")
async def demo_result(subject: str = "C"):
    demos = {
        "C": {"subject":"C","subject_name":"货币资金","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"货币资金期末余额6,055,000元。①工行渝北支行存在银收企未记3,500元、银付企未记1,200元，调节后余额一致；②其他货币资金-保证金300,000元需函证确认；③现金盘点差异-900元需查明。",
              "sheets":["C-1 货币资金审定表","C-2 货币资金明细表","C-3 现金监盘表","C-4 银行余额调节表检查","C-5 银行函证汇总表","C-6 大额双向核对表","C-7 银行存款利息测算表"],
              "audit_findings":[{"severity":"warning","msg":"工行渝北支行存在未达账项（银收企未记3,500元、银付企未记1,200元）"},{"severity":"error","msg":"建行、招行等账户未取得银行对账单，无法执行余额调节"},{"severity":"warning","msg":"库存现金盘点差异-900元"},{"severity":"warning","msg":"其他货币资金-保证金300,000元未函证"},{"severity":"info","msg":"12月31日大额转账至建行500,000元，关注期后资金流向"}]},
        "D": {"subject":"D","subject_name":"应收票据","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"应收票据期末余额2,700,000元（银承1,800,000元、商承900,000元）。商承PJ-007（135,000元）已逾期未兑付，需转入应收账款并评估减值。已贴现银承620,000元终止确认合规。",
              "sheets":["D-1 应收票据审定表","D-2 应收票据明细表(按客户)","D-3 应收票据盘点表","D-4 票据备查簿核对表","D-5 应收票据贴现检查表","D-6 预期信用损失检查表","D-7 坏账准备计算表"],
              "audit_findings":[{"severity":"error","msg":"商承PJ-007（135,000元）已逾期，需转入应收账款并减值测试"},{"severity":"warning","msg":"商承PJ-003（265,000元）信用等级较低，关注兑付风险"},{"severity":"warning","msg":"已背书银承PJ-006（850,000元）未到期，需检查终止确认条件"},{"severity":"info","msg":"已贴现银承620,000元，会计处理符合终止确认条件"}]},
        "E": {"subject":"E","subject_name":"其他应收款","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"其他应收款期末余额965,000元。股东王五借款500,000元属关联方资金占用，需检查审批、计息及披露。前员工李四借款30,000元挂账25个月，应全额计提坏账。重庆鑫源装备保证金50,000元（关联方）需披露。",
              "sheets":["E-1 其他应收款审定表","E-2 其他应收款明细表","E-3 账龄分析表","E-4 函证汇总表","E-5 控股股东资金占用审核表","E-6 关联方交易核对表","E-7 预期信用损失检查表"],
              "audit_findings":[{"severity":"error","msg":"股东王五借款500,000元（关联方资金占用），需检查审批和披露"},{"severity":"error","msg":"前员工李四借款30,000元挂账25个月，应全额计提坏账"},{"severity":"warning","msg":"重庆鑫源装备保证金50,000元（关联方），需在附注中披露"},{"severity":"warning","msg":"四川恒力机械资金拆借30万元，需检查借款协议及利息条款"}]},
        "F": {"subject":"F","subject_name":"预付账款","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"预付账款期末余额2,660,000元。华为技术预付850,000元（100%付款）货物未到，关注合理性。重庆渝快物流180,000元挂账14个月，关注无法履约风险。",
              "sheets":["F-1 预付账款审定表","F-2 预付账款明细表","F-3 账龄分析表","F-4 函证汇总表","F-5 长期挂账查验表","F-6 减值准备复核表"],
              "audit_findings":[{"severity":"error","msg":"华为技术850,000元全额预付未到货，关注商业合理性和合同条款"},{"severity":"error","msg":"重庆渝快物流180,000元挂账14个月，关注是否无法履约"},{"severity":"warning","msg":"重庆钢铁设备480,000元（付款40%），需核对到货比例冲减"},{"severity":"info","msg":"京东企业购已到货结算，处理正常"}]},
        "G": {"subject":"G","subject_name":"预收账款","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"预收账款期末余额5,075,000元。赛力斯集团1,800,000元（100%收款）未交付，关注推迟确认收入风险。长安福特420,000元挂账15个月需查明原因。",
              "sheets":["G-1 预收账款审定表","G-2 预收账款明细表","G-3 账龄分析表","G-4 函证汇总表","G-5 收入确认检查表","G-6 长期挂账查验表"],
              "audit_findings":[{"severity":"error","msg":"赛力斯集团1,800,000元全额预收未交付，关注推迟确认收入风险"},{"severity":"error","msg":"长安福特420,000元挂账15个月，关注合同纠纷或应退款"},{"severity":"warning","msg":"重庆长安汽车1,250,000元（收款50%）部分交付，按履约进度确认收入"},{"severity":"warning","msg":"重庆燃气320,000元分批交付中，关注收入确认时点"}]},
        "H": {"subject":"H","subject_name":"短期借款","entity_name":"重庆蛮先进智能制造有限公司","period":"2024-12-31",
              "summary":"短期借款期末余额19,500,000元。重庆农商行利息差异3,000元需补提。招行重庆分行2,000,000元已到期需确认续贷。抵押借款5,000,000元需检查抵押合同。",
              "sheets":["H-1 短期借款审定表","H-2 借款明细及利息复核表","H-3 函证汇总表","H-4 贷款卡核对表","H-5 逾期贷款检查表","H-6 抵押质押担保统计表"],
              "audit_findings":[{"severity":"error","msg":"重庆农商行利息差异3,000元（应计23,000 vs 账面20,000），需补提"},{"severity":"warning","msg":"招行重庆分行2,000,000元已到期，需确认续贷或还款计划"},{"severity":"warning","msg":"抵押借款5,000,000元（厂房+设备），需检查抵押合同及所有权"},{"severity":"info","msg":"中行重庆分行8,000,000元新借款，利息差异333元（不重大）"}]},
    }
    if subject not in demos: raise HTTPException(404, f"未知科目: {subject}")
    result = dict(demos[subject]); result["success"] = True
    # 优先使用实际生成的底稿和日志
    DEMO_DIR = BASE_DIR / "demo_outputs"
    actual_xlsx = list(DEMO_DIR.glob(f"{subject}_*审计底稿*.xlsx")) if DEMO_DIR.exists() else []
    actual_logs = list(DEMO_DIR.glob(f"{subject}_*审计日志*.md")) if DEMO_DIR.exists() else []
    if actual_xlsx:
        actual_xlsx.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        result["output_file"] = str(actual_xlsx[0]); result["is_actual"] = True
    else:
        result["output_file"] = ""; result["is_actual"] = False
    if actual_logs:
        actual_logs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
        result["log_file"] = str(actual_logs[0])
        result["log_content"] = actual_logs[0].read_text(encoding="utf-8")[:3000] if actual_logs[0].exists() else ""
    else:
        result["log_file"] = ""; result["log_content"] = ""
    result["risk_count"] = len(result.get("audit_findings", []))
    return result

@app.get("/api/demo-files")
async def demo_files(subject: str = "C"):
    SAMPLE_DIR = BASE_DIR / "samples"
    fm = {"C":["C_科目余额表.xlsx","C_明细账.xlsx","C_银行对账单.xlsx","C_现金盘点表.xlsx"],
          "D":["D_科目余额表.xlsx","D_应收票据备查簿.xlsx"],
          "E":["E_科目余额表.xlsx","E_其他应收款明细表.xlsx"],
          "F":["F_科目余额表.xlsx","F_预付账款明细表.xlsx"],
          "G":["G_科目余额表.xlsx","G_预收账款明细表.xlsx"],
          "H":["H_科目余额表.xlsx","H_短期借款台账.xlsx"]}
    files = fm.get(subject, [])
    paths = [str(SAMPLE_DIR/f) for f in files if (SAMPLE_DIR/f).exists()]
    return {"success": True, "files": paths, "count": len(paths)}

@app.get("/api/download")
async def download(path: str):
    if not os.path.exists(path): raise HTTPException(404, "文件不存在")
    return FileResponse(path, filename=os.path.basename(path),
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ==================== 用户面板 API ====================
TEMPLATE_DIR = BASE_DIR / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)

@app.post("/api/templates/upload")
async def upload_template(subject: str = "C", file: UploadFile = None):
    """上传某科目的底稿模板"""
    if not file or not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(400, "请上传Excel模板文件")
    content = await file.read()
    safe_name = re.sub(r'[\\/*?:<>|]', '_', file.filename)
    path = TEMPLATE_DIR / f"{subject}_{safe_name}"
    path.write_bytes(content)
    return {"success": True, "path": str(path), "subject": subject}

@app.get("/api/templates")
async def list_templates():
    """列出所有已上传的模板"""
    templates = []
    for f in TEMPLATE_DIR.glob("*.xlsx"):
        subject = f.name.split('_')[0] if '_' in f.name else '?'
        templates.append({"subject": subject, "name": f.name, "path": str(f), "size": f.stat().st_size, "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M")})
    return {"templates": sorted(templates, key=lambda x: x['modified'], reverse=True)}

@app.get("/api/recent")
async def recent_records():
    """获取近期生成的审计记录"""
    records = []
    # 底稿文件
    for f in sorted(OUTPUT_DIR.glob("*.xlsx"), key=lambda x: x.stat().st_mtime, reverse=True)[:20]:
        records.append({"type": "底稿", "name": f.name, "path": str(f), "size": f.stat().st_size, "time": datetime.fromtimestamp(f.stat().st_mtime).strftime("%m-%d %H:%M")})
    # 日志文件
    for f in sorted(OUTPUT_DIR.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True)[:20]:
        records.append({"type": "日志", "name": f.name, "path": str(f), "size": f.stat().st_size, "time": datetime.fromtimestamp(f.stat().st_mtime).strftime("%m-%d %H:%M")})
    # 合并排序
    records.sort(key=lambda x: x['time'], reverse=True)
    return {"records": records[:30]}

@app.get("/api/risk-summary")
async def risk_summary():
    """返回近期风险摘要"""
    risks = []
    # 扫描demo_outputs中的日志
    DEMO_DIR = BASE_DIR / "demo_outputs"
    for f in sorted(DEMO_DIR.glob("*审计日志*"), key=lambda x: x.stat().st_mtime, reverse=True)[:6]:
        content = f.read_text(encoding="utf-8")[:3000]
        # 提取风险相关行
        for line in content.split('\n'):
            if any(kw in line for kw in ['🔴','🟡','风险','异常','需关注','差异','逾期','未回函']):
                risks.append({"source": f.name.replace('_审计日志.md','').replace('_审计日志.docx',''), "finding": line.strip('| -*'), "time": datetime.fromtimestamp(f.stat().st_mtime).strftime("%m-%d %H:%M")})
    return {"risks": risks[:20]}


@app.get("/api/health")
async def health():
    return {"status":"ok","model":ANTHROPIC_MODEL,"subjects":list(SUBJECT_MAP.keys())}

@app.get("/")
async def root():
    return FileResponse(BASE_DIR / "static" / "index.html")

if __name__ == "__main__":
    import uvicorn; uvicorn.run(app, host="127.0.0.1", port=8800)
